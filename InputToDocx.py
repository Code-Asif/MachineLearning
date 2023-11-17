from docx import Document

def create_preformatted_document(user_input):
    doc = Document()

    doc.add_heading('Preformatted Document', level=1)

    doc.add_paragraph(f"User input: {user_input}")

    doc.save('preformatted_document.docx')

if __name__ == "__main__":
    user_input = input("Enter text to be inserted into the document: ")

    create_preformatted_document(user_input)

    print("Document created and saved as 'preformatted_document.docx'")
