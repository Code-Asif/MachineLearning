from docx import Document
from docx.shared import Inches
from PIL import Image

def convert_image_to_word(image_path, output_docx_path):
    # Open the image
    img = Image.open(image_path)

    # Create a new Word document
    doc = Document()

    # Add a title to the document
    doc.add_heading('Image to Word Document', level=1)

    # Add the image to the Word document
    doc.add_picture(image_path, width=Inches(4))  # You can adjust the width as needed

    # Save the Word document
    doc.save(output_docx_path)

    print(f"Conversion successful. Word document saved at {output_docx_path}")

if __name__ == "__main__":
    # Replace 'path/to/your/image.jpg' with the path to your image file
    image_path = 'Photo.jpg'

    # Replace 'output_document.docx' with the desired output Word document path
    output_docx_path = 'output_document.docx'

    convert_image_to_word(image_path, output_docx_path)
